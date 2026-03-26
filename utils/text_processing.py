import fitz  # PyMuPDF
import docx
import io

def extract_text(filepath_or_bytes, file_type):
    """
    Extract text from PDF, DOCX, or TXT.
    `filepath_or_bytes` can be a string path or a bytes object.
    `file_type` should be 'pdf', 'docx', or 'txt'.
    """
    file_type = file_type.lower()
    text = ""
    
    try:
        if file_type == 'pdf':
            if isinstance(filepath_or_bytes, bytes):
                doc = fitz.open(stream=filepath_or_bytes, filetype="pdf")
            else:
                doc = fitz.open(filepath_or_bytes)
            for page in doc:
                text += page.get_text() + "\n"
        elif file_type == 'docx':
            if isinstance(filepath_or_bytes, bytes):
                doc = docx.Document(io.BytesIO(filepath_or_bytes))
            else:
                doc = docx.Document(filepath_or_bytes)
            text = "\n".join([para.text for para in doc.paragraphs])
        elif file_type == 'txt':
            if isinstance(filepath_or_bytes, bytes):
                text = filepath_or_bytes.decode('utf-8', errors='ignore')
            else:
                with open(filepath_or_bytes, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
    except Exception as e:
        print(f"Error extracting text: {e}")
        return ""
        
    # Basic cleanup
    text = " ".join(text.split())
    return text

def chunk_text(text, chunk_size=400, overlap=50):
    """
    Basic character-level splitting by words to approximate tokens.
    """
    words = text.split()
    chunks = []
    
    if len(words) == 0:
        return chunks
        
    i = 0
    while i < len(words):
        chunk = " ".join(words[i:i + chunk_size])
        chunks.append(chunk)
        i += chunk_size - overlap
        
    return chunks
